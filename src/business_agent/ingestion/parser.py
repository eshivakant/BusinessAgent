from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import httpx
from docx import Document
from pypdf import PdfReader

SUPPORTED_SOURCE_TYPES = {"txt", "pdf", "docx", "png", "jpg", "jpeg", "gif", "webp"}


@dataclass(frozen=True)
class ParsedDocument:
    source_uri: str
    source_type: str
    text: str


def load_document_from_uri(source_uri: str, allowed_local_dir: str | None = None) -> ParsedDocument:
    parsed_uri = urlparse(source_uri)
    if parsed_uri.scheme in {"http", "https"}:
        response = httpx.get(source_uri, timeout=30.0, follow_redirects=True)
        response.raise_for_status()
        source_type = resolve_source_type(source_uri, response.headers.get("Content-Type"))
        content = response.content
    else:
        path = Path(source_uri).expanduser().resolve()
        if not path.exists() or not path.is_file():
            raise ValueError(f"Document does not exist: {source_uri}")
        if allowed_local_dir:
            base_dir = Path(allowed_local_dir).expanduser().resolve()
            if path != base_dir and base_dir not in path.parents:
                raise PermissionError(
                    "Local file is outside INGESTION_ALLOWED_LOCAL_DIR; document read denied."
                )
        source_type = resolve_source_type(path.name, None)
        content = path.read_bytes()

    text = parse_document_bytes(content, source_type)
    normalized = normalize_text(text)
    if not normalized:
        raise ValueError("The document has no extractable text.")

    return ParsedDocument(source_uri=source_uri, source_type=source_type, text=normalized)


def resolve_source_type(source_uri: str, content_type: str | None) -> str:
    suffix = Path(urlparse(source_uri).path or source_uri).suffix.lower()
    if suffix == ".txt":
        return "txt"
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
        return suffix.lstrip(".")

    if content_type:
        lowered = content_type.lower()
        if "text/plain" in lowered:
            return "txt"
        if "application/pdf" in lowered:
            return "pdf"
        if "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in lowered:
            return "docx"
        if "image/png" in lowered:
            return "png"
        if "image/jpeg" in lowered:
            return "jpg"
        if "image/gif" in lowered:
            return "gif"
        if "image/webp" in lowered:
            return "webp"

    raise ValueError("Unsupported source type. Use txt, pdf, docx, or image (png/jpg/jpeg/gif/webp).")


def parse_document_bytes(content: bytes, source_type: str) -> str:
    if source_type == "txt":
        return content.decode("utf-8", errors="ignore")
    if source_type == "pdf":
        reader = PdfReader(BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if source_type == "docx":
        document = Document(BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    if source_type in {"png", "jpg", "jpeg", "gif", "webp"}:
        return f"[IMAGE: {source_type.upper()} - OCR needed]"
    raise ValueError(f"Unsupported source type: {source_type}")


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)

