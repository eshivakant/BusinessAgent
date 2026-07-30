from __future__ import annotations

import re
import shutil
import subprocess
import uuid
from datetime import date, datetime, time, timezone
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument

try:
    from docxtpl import DocxTemplate
except ImportError:  # pragma: no cover - import fallback for test environments
    DocxTemplate = None

from business_agent.ingestion.chunking import chunk_text
from business_agent.ingestion.parser import load_document_from_uri
from business_agent.ingestion.summarizer import Summarizer
from business_agent.memory.models import MemoryPayload, MemoryRecord
from business_agent.memory.store import MemoryStore
from business_agent.property.models import Property, Tenant
from business_agent.property.registry import PropertyRegistry
from business_agent.tenancy.models import GeneratedAgreement, TemplateSelectionResult, TenantDocument
from business_agent.tenancy.registry import TenancyRegistry


_PLACEHOLDER_NAMES = [
    "TENANT_FULL_NAME",
    "TENANT_EMAIL",
    "PROPERTY_ADDRESS",
    "PROPERTY_POSTCODE",
    "MONTHLY_RENT",
    "LEASE_START_DATE",
    "LEASE_END_DATE",
    "DEPOSIT_AMOUNT",
    "LANDLORD_NAME",
    "LANDLORD_ADDRESS",
]


class TenancyService:
    def __init__(
        self,
        tenancy_registry: TenancyRegistry,
        property_registry: PropertyRegistry,
        memory_store: MemoryStore | None = None,
        summarizer: Summarizer | None = None,
        chunk_size: int = 1200,
        chunk_overlap: int = 200,
        max_document_chars: int = 200000,
        storage_dir: str = "/data/tenant-docs",
        template_dir: str = "/data/agreement-templates",
        generated_dir: str = "/data/generated-agreements",
        llm_client: Any | None = None,
        allowed_local_dir: str | None = None,
    ) -> None:
        self._tenancy_registry = tenancy_registry
        self._property_registry = property_registry
        self._memory_store = memory_store
        self._summarizer = summarizer
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
        self._max_document_chars = max_document_chars
        self._storage_dir = Path(storage_dir)
        self._template_dir = Path(template_dir)
        self._generated_dir = Path(generated_dir)
        self._llm_client = llm_client
        self._allowed_local_dir = Path(allowed_local_dir) if allowed_local_dir else self._storage_dir

    def create_tenancy(
        self,
        property_id: str,
        full_name: str,
        *,
        email: str | None = None,
        phone: str | None = None,
        lease_start: date | None = None,
        lease_end: date | None = None,
        monthly_rent: Decimal | None = None,
        deposit: Decimal | None = None,
        notes: str | None = None,
    ) -> Tenant:
        if not full_name.strip():
            raise ValueError("full_name is required")
        if self._property_registry.get_property(property_id) is None:
            raise ValueError(f"Property not found: {property_id}")

        tenant = Tenant(
            id=f"tenancy-{uuid.uuid4().hex[:8]}",
            property_id=property_id,
            name=full_name.strip(),
            email=email,
            phone=phone,
            lease_start=lease_start or date.today(),
            lease_end=lease_end or date.today(),
            monthly_rent=monthly_rent or Decimal("0"),
            deposit=deposit or Decimal("0"),
            notes=notes,
            is_active=True,
            created_at=datetime.now(timezone.utc),
            updated_at=datetime.now(timezone.utc),
            full_name=full_name.strip(),
        )
        self._tenancy_registry.create_tenancy(tenant)
        return tenant

    def get_tenancy(self, tenancy_id: str) -> Tenant | None:
        return self._tenancy_registry.get_tenancy(tenancy_id)

    def list_tenancies(self, property_id: str | None = None, active_only: bool = True) -> list[Tenant]:
        return self._tenancy_registry.list_tenancies(property_id=property_id, active_only=active_only)

    def update_tenancy(self, tenancy_id: str, updates: dict[str, Any]) -> Tenant | None:
        tenancy = self._tenancy_registry.get_tenancy(tenancy_id)
        if tenancy is None:
            return None

        for field_name in [
            "full_name",
            "email",
            "phone",
            "lease_start",
            "lease_end",
            "monthly_rent",
            "deposit",
            "is_active",
            "notes",
            "date_of_birth",
            "national_insurance_number",
            "passport_number",
            "employer_name",
            "annual_income",
        ]:
            if field_name in updates and updates[field_name] is not None:
                setattr(tenancy, field_name, updates[field_name])
                if field_name == "full_name" and not getattr(tenancy, "name", None):
                    tenancy.name = updates[field_name]

        tenancy.updated_at = datetime.now(timezone.utc)
        self._tenancy_registry.update_tenancy(tenancy)
        return tenancy

    def store_document(self, tenancy_id: str, source_path: str | Path, filename: str | None = None, event_date: date | None = None) -> TenantDocument:
        tenancy = self._tenancy_registry.get_tenancy(tenancy_id)
        if tenancy is None:
            raise ValueError(f"Tenancy not found: {tenancy_id}")

        source_path = Path(source_path)
        if not source_path.exists():
            raise ValueError(f"Document does not exist: {source_path}")

        destination_dir = self._storage_dir / tenancy_id
        destination_dir.mkdir(parents=True, exist_ok=True)
        suffix = source_path.suffix.lower() or ".bin"
        document_id = uuid.uuid4().hex
        stored_path = destination_dir / f"{document_id}{suffix}"
        shutil.copy2(source_path, stored_path)

        allowed_local_dir = self._storage_dir
        parsed = load_document_from_uri(str(stored_path), allowed_local_dir=str(allowed_local_dir))
        text = parsed.text[: self._max_document_chars]
        summary = self._summarizer.summarize(text) if self._summarizer else text[:400]
        chunks = chunk_text(text, chunk_size=self._chunk_size, overlap=self._chunk_overlap)

        ingested_at = datetime.now(timezone.utc)
        effective_date = self._compute_effective_date(event_date or self._extract_event_date(text), ingested_at)
        extracted_fields = self.extract_structured_fields(text)

        qdrant_ids: list[str] = []
        if self._memory_store is not None:
            summary_payload = MemoryPayload(
                event_date=event_date or self._extract_event_date(text),
                ingested_at=ingested_at,
                effective_date=effective_date,
                source_type="tenant_document",
                source_uri=str(stored_path),
                archived_file_path=None,
                record_type="summary",
                chunk_index=None,
                chunk_count=len(chunks),
                summary=summary,
                property_address=self._property_registry.get_property(tenancy.property_id).address if self._property_registry.get_property(tenancy.property_id) else None,
                property_id=tenancy.property_id,
                tenancy_id=tenancy.id,
                document_id=document_id,
                document_type=parsed.source_type,
                amount=None,
            )
            summary_id = f"{document_id}:summary"
            qdrant_ids.append(summary_id)
            records = [MemoryRecord(id=summary_id, text=summary, payload=summary_payload)]
            for index, chunk in enumerate(chunks):
                payload = MemoryPayload(
                    event_date=event_date or self._extract_event_date(text),
                    ingested_at=ingested_at,
                    effective_date=effective_date,
                    source_type="tenant_document",
                    source_uri=str(stored_path),
                    archived_file_path=None,
                    record_type="chunk",
                    chunk_index=index,
                    chunk_count=len(chunks),
                    summary=summary,
                    property_address=self._property_registry.get_property(tenancy.property_id).address if self._property_registry.get_property(tenancy.property_id) else None,
                    property_id=tenancy.property_id,
                    tenancy_id=tenancy.id,
                    document_id=document_id,
                    document_type=parsed.source_type,
                    amount=None,
                )
                chunk_id = f"{document_id}:chunk:{index}"
                qdrant_ids.append(chunk_id)
                records.append(MemoryRecord(id=chunk_id, text=chunk, payload=payload))
            self._memory_store.upsert(records)

        document = TenantDocument(
            id=document_id,
            tenancy_id=tenancy.id,
            filename=filename or source_path.name,
            stored_path=str(stored_path),
            document_type=parsed.source_type,
            ingested_at=ingested_at,
            extracted_fields=extracted_fields,
            qdrant_ids=qdrant_ids,
            property_id=tenancy.property_id,
            source_uri=str(stored_path),
            summary=summary,
            chunk_count=len(chunks),
        )
        self._tenancy_registry.add_document(document)

        if extracted_fields:
            for field_name, value in extracted_fields.items():
                if not getattr(tenancy, field_name, None):
                    setattr(tenancy, field_name, value)
            tenancy.updated_at = datetime.now(timezone.utc)
            self._tenancy_registry.update_tenancy(tenancy)

        return document

    def search_documents(self, query: str, *, tenancy_id: str | None = None, date_from: date | None = None, date_to: date | None = None, top_k: int = 5) -> list[Any]:
        if self._memory_store is None:
            return []

        from business_agent.memory.models import MemoryQueryInput

        request = MemoryQueryInput(
            query=query,
            top_k=top_k,
            date_from=date_from,
            date_to=date_to,
            source_type="tenant_document",
        )
        matches = self._memory_store.query(request)
        if tenancy_id is None:
            return matches
        return [match for match in matches if match.payload.tenancy_id == tenancy_id]

    def list_documents(self, tenancy_id: str) -> list[TenantDocument]:
        return self._tenancy_registry.list_documents(tenancy_id)

    def discover_templates(self) -> list[Path]:
        if not self._template_dir.exists():
            return []
        return sorted([path for path in self._template_dir.glob("*.docx")])

    def select_template(self, template_hint: str | None = None) -> TemplateSelectionResult:
        templates = self.discover_templates()
        if not templates:
            return TemplateSelectionResult(template_name=None, candidates=[], needs_selection=False)

        if template_hint:
            hint = template_hint.strip().lower()
            exact = [path.stem for path in templates if path.stem.lower() == hint]
            if exact:
                return TemplateSelectionResult(template_name=exact[0], candidates=exact, needs_selection=False)
            fuzzy = [path.stem for path in templates if hint in path.stem.lower()]
            if fuzzy:
                return TemplateSelectionResult(template_name=fuzzy[0] if len(fuzzy) == 1 else None, candidates=fuzzy, needs_selection=len(fuzzy) > 1)
            return TemplateSelectionResult(template_name=None, candidates=[path.stem for path in templates], needs_selection=True)

        if len(templates) == 1:
            return TemplateSelectionResult(template_name=templates[0].stem, candidates=[templates[0].stem], needs_selection=False)
        return TemplateSelectionResult(template_name=None, candidates=[path.stem for path in templates], needs_selection=True)

    def generate_agreement(
        self,
        tenancy_id: str,
        *,
        template_name: str | None = None,
        values: dict[str, Any] | None = None,
        missing_values: dict[str, Any] | None = None,
    ) -> tuple[GeneratedAgreement, list[str]]:
        tenancy = self._tenancy_registry.get_tenancy(tenancy_id)
        if tenancy is None:
            raise ValueError(f"Tenancy not found: {tenancy_id}")

        property_record = self._property_registry.get_property(tenancy.property_id)
        templates = self.discover_templates()
        if not templates:
            raise ValueError("No agreement templates are available")

        resolved_name = template_name
        if not resolved_name:
            selection = self.select_template()
            if selection.needs_selection or selection.template_name is None:
                raise ValueError("Multiple templates matched; please select one")
            resolved_name = selection.template_name

        if resolved_name:
            matching_template = next((path for path in templates if path.stem.lower() == resolved_name.lower()), None)
            if matching_template is None:
                raise ValueError(f"Template not found: {resolved_name}")
        else:
            matching_template = templates[0]

        context = self._build_context(tenancy, property_record)
        if values:
            context.update(values)
        if missing_values:
            context.update(missing_values)

        output_dir = self._generated_dir / tenancy_id
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"agreement_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}.docx"
        self._render_template(matching_template, output_path, context)

        agreement = GeneratedAgreement(
            id=f"agreement-{uuid.uuid4().hex[:8]}",
            tenancy_id=tenancy.id,
            template_name=matching_template.stem,
            generated_at=datetime.now(timezone.utc),
            stored_path=str(output_path),
            pdf_path=None,
        )
        self._tenancy_registry.create_agreement(agreement)
        unresolved = self._find_unresolved_placeholders(output_path)
        return agreement, unresolved

    def convert_to_pdf(self, agreement_id: str) -> str:
        agreement = self._tenancy_registry.get_agreement(agreement_id)
        if agreement is None:
            raise ValueError(f"Agreement not found: {agreement_id}")
        source_path = Path(agreement.stored_path)
        if not source_path.exists():
            raise ValueError(f"Agreement document not found: {source_path}")

        pdf_path = source_path.with_suffix(".pdf")
        if shutil.which("libreoffice"):
            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    str(source_path),
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            return str(pdf_path)

        raise RuntimeError("No PDF conversion backend is available")

    def extract_structured_fields(self, text: str) -> dict[str, Any]:
        fields: dict[str, Any] = {}

        for pattern, field_name in [
            (r"\b(?:full name|name|tenant)[:\s]+([A-Z][a-z]+(?:[ ]+[A-Z][a-z]+)+)", "full_name"),
            (r"\b(?:date of birth|dob)[:\s]+(\d{4}-\d{2}-\d{2})", "date_of_birth"),
            (r"\b(?:national insurance|ni)[:\s]+([A-Z]{2}(?:\s?\d{2}){3}(?:\s?[A-Z])?)", "national_insurance_number"),
            (r"\b(?:passport|passport number)[:\s]+([A-Z0-9]{4,12})", "passport_number"),
            (r"\b(?:email|e-mail)[:\s]+([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})", "email"),
            (r"\b(?:phone|mobile|tel)[:\s]+([+0-9\s-]{5,})", "phone"),
            (r"\b(?:employer|company)[:\s]+([A-Za-z0-9&.,' -]{2,})", "employer_name"),
            (r"\b(?:annual income|income)[:\s]+£?([0-9,]+(?:\.\d{2})?)", "annual_income"),
        ]:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                if field_name == "annual_income":
                    value = Decimal(value.replace(",", ""))
                elif field_name == "date_of_birth":
                    value = date.fromisoformat(value)
                fields[field_name] = value

        return fields

    def _build_context(self, tenancy: Tenant, property_record: Property | None) -> dict[str, Any]:
        values: dict[str, Any] = {}
        values["TENANT_FULL_NAME"] = tenancy.full_name or tenancy.name
        values["TENANT_EMAIL"] = tenancy.email or ""
        values["PROPERTY_ADDRESS"] = property_record.address if property_record else ""
        values["PROPERTY_POSTCODE"] = property_record.postcode if property_record else ""
        values["MONTHLY_RENT"] = self._format_currency(tenancy.monthly_rent)
        values["LEASE_START_DATE"] = self._format_date(tenancy.lease_start)
        values["LEASE_END_DATE"] = self._format_date(tenancy.lease_end)
        values["DEPOSIT_AMOUNT"] = self._format_currency(tenancy.deposit)
        values["LANDLORD_NAME"] = "Landlord"
        values["LANDLORD_ADDRESS"] = ""
        return {key: values.get(key, "") for key in _PLACEHOLDER_NAMES}

    def _render_template(self, template_path: Path, output_path: Path, context: dict[str, Any]) -> None:
        if DocxTemplate is not None:
            template = DocxTemplate(str(template_path))
            template.render(context)
            template.save(str(output_path))
            self._replace_placeholders_in_docx(output_path, context)
            return

        document = DocxDocument(str(template_path))
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, context)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, context)
        document.save(str(output_path))

    def _replace_placeholders_in_docx(self, output_path: Path, context: dict[str, Any]) -> None:
        document = DocxDocument(str(output_path))
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, context)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, context)
        document.save(str(output_path))

    def _replace_in_paragraph(self, paragraph: Any, context: dict[str, Any]) -> None:
        if not paragraph.text:
            return
        text = paragraph.text
        for placeholder, value in context.items():
            text = text.replace(f"*|{placeholder}|*", str(value))
        paragraph.text = text

    def _find_unresolved_placeholders(self, output_path: Path) -> list[str]:
        document = DocxDocument(str(output_path))
        unresolved: list[str] = []
        for paragraph in document.paragraphs:
            if "*|" in paragraph.text:
                unresolved.extend(re.findall(r"\*\|([A-Z_]+)\|\*", paragraph.text))
        return sorted(set(unresolved))

    def _compute_effective_date(self, event_date: date | None, ingested_at: datetime) -> datetime:
        if event_date is None:
            return ingested_at
        return datetime.combine(event_date, time.min, tzinfo=timezone.utc)

    def _extract_event_date(self, text: str) -> date | None:
        patterns = [
            r"(\d{4}-\d{2}-\d{2})",
            r"(\d{1,2}\s+[A-Za-z]+\s+\d{4})",
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                try:
                    return date.fromisoformat(match.group(1))
                except ValueError:
                    try:
                        return datetime.strptime(match.group(1), "%d %B %Y").date()
                    except ValueError:
                        continue
        return None

    def _format_date(self, value: date | None) -> str:
        if value is None:
            return ""
        return value.strftime("%d %B %Y")

    def _format_currency(self, value: Decimal | None) -> str:
        if value is None:
            return ""
        return f"£{value:,.2f}"
