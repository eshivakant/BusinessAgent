from __future__ import annotations

from datetime import date, datetime
from decimal import Decimal
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from business_agent.property.models import Property, PropertyStatus
from business_agent.property.registry import InMemoryPropertyRegistry
from business_agent.tenancy.registry import InMemoryTenancyRegistry
from business_agent.tenancy.service import TenancyService


@pytest.mark.unit
def test_tenancy_context_mapping_and_template_selection(tmp_path: Path) -> None:
    property_registry = InMemoryPropertyRegistry()
    tenancy_registry = InMemoryTenancyRegistry()
    service = TenancyService(
        tenancy_registry=tenancy_registry,
        property_registry=property_registry,
        memory_store=None,
        summarizer=None,
        storage_dir=str(tmp_path / "tenant-docs"),
        template_dir=str(tmp_path / "templates"),
        generated_dir=str(tmp_path / "generated"),
        allowed_local_dir=str(tmp_path / "docs"),
    )

    property_record = Property(
        id="prop-1",
        address="133 Example Street",
        purchase_date=None,
        purchase_price=None,
        current_value=None,
        status=PropertyStatus.OWNED,
        postcode="SW1A 1AA",
    )
    property_registry.add_property(property_record)

    tenancy = service.create_tenancy(
        property_id=property_record.id,
        full_name="Ada Lovelace",
        lease_start=date(2024, 1, 1),
        lease_end=date(2025, 1, 1),
        monthly_rent=Decimal("1200.00"),
        deposit=Decimal("3000.00"),
    )

    context = service._build_context(tenancy, property_record)
    assert context["TENANT_FULL_NAME"] == "Ada Lovelace"
    assert context["PROPERTY_ADDRESS"] == "133 Example Street"
    assert context["MONTHLY_RENT"] == "£1,200.00"
    assert context["PROPERTY_POSTCODE"] == "SW1A 1AA"

    templates_dir = Path(service._template_dir)
    templates_dir.mkdir(parents=True, exist_ok=True)
    for filename in ["residential_lease.docx", "commercial_lease.docx"]:
        document = Document()
        document.add_paragraph("*|TENANT_FULL_NAME|*")
        document.save(templates_dir / filename)

    selection = service.select_template("lease")
    assert selection.needs_selection is True
    assert set(selection.candidates) == {"commercial_lease", "residential_lease"}


@pytest.mark.unit
def test_tenancy_structured_field_extraction() -> None:
    property_registry = InMemoryPropertyRegistry()
    tenancy_registry = InMemoryTenancyRegistry()
    service = TenancyService(
        tenancy_registry=tenancy_registry,
        property_registry=property_registry,
        memory_store=None,
        summarizer=None,
    )

    text = (
        "Tenant: Grace Hopper\n"
        "Date of Birth: 1990-01-02\n"
        "Email: grace@example.com\n"
        "Phone: +44 7711 223344\n"
        "Annual Income: £45,000\n"
        "National Insurance: QQ 12 34 56 A"
    )
    fields = service.extract_structured_fields(text)

    assert fields["full_name"] == "Grace Hopper"
    assert fields["date_of_birth"] == date(1990, 1, 2)
    assert fields["email"] == "grace@example.com"
    assert fields["annual_income"] == Decimal("45000")
    assert fields["national_insurance_number"] == "QQ 12 34 56 A"


@pytest.mark.unit
def test_tenancy_document_storage_path_is_scoped_to_tenancy(tmp_path: Path) -> None:
    property_registry = InMemoryPropertyRegistry()
    tenancy_registry = InMemoryTenancyRegistry()
    service = TenancyService(
        tenancy_registry=tenancy_registry,
        property_registry=property_registry,
        memory_store=None,
        summarizer=None,
        storage_dir=str(tmp_path / "tenant-docs"),
        template_dir=str(tmp_path / "templates"),
        generated_dir=str(tmp_path / "generated"),
        allowed_local_dir=str(tmp_path / "docs"),
    )
    property_registry.add_property(
        Property(
            id="prop-2",
            address="12 River Lane",
            purchase_date=None,
            purchase_price=None,
            current_value=None,
            status=PropertyStatus.OWNED,
        )
    )

    tenancy = service.create_tenancy(
        property_id="prop-2",
        full_name="Linus Torvalds",
        lease_start=date(2024, 1, 1),
        lease_end=date(2026, 1, 1),
        monthly_rent=Decimal("900.00"),
        deposit=Decimal("1000.00"),
    )

    source_path = tmp_path / "upload.txt"
    source_path.write_text("A tenant document for testing", encoding="utf-8")
    document = service.store_document(tenancy.id, source_path, filename="upload.txt")

    stored_path = Path(document.stored_path)
    assert stored_path.parent.name == tenancy.id
    assert stored_path.exists()
    assert stored_path.suffix == ".txt"


@pytest.mark.unit
def test_tenancy_updates_only_non_none_values() -> None:
    property_registry = InMemoryPropertyRegistry()
    tenancy_registry = InMemoryTenancyRegistry()
    service = TenancyService(
        tenancy_registry=tenancy_registry,
        property_registry=property_registry,
        memory_store=None,
        summarizer=None,
    )
    property_registry.add_property(
        Property(
            id="prop-3",
            address="9 Willow Road",
            purchase_date=None,
            purchase_price=None,
            current_value=None,
            status=PropertyStatus.OWNED,
        )
    )
    tenancy = service.create_tenancy(
        property_id="prop-3",
        full_name="Nina Simone",
        email="nina@example.com",
        phone="0200 111 2222",
        lease_start=date(2024, 1, 1),
        lease_end=date(2025, 1, 1),
        monthly_rent=Decimal("1000.00"),
        deposit=Decimal("2000.00"),
    )

    updated = service.update_tenancy(tenancy.id, {"email": "new@example.com", "phone": None, "monthly_rent": Decimal("1100.00")})
    assert updated is not None
    assert updated.email == "new@example.com"
    assert updated.phone == "0200 111 2222"
    assert updated.monthly_rent == Decimal("1100.00")


@pytest.mark.e2e
def test_tenant_crud_and_telegram_flow(fast_e2e_harness: Any) -> None:
    property_record = Property(
        id="prop-tenant-1",
        address="101 High Street",
        purchase_date=None,
        purchase_price=None,
        current_value=None,
        status=PropertyStatus.OWNED,
    )
    fast_e2e_harness.property_registry.add_property(property_record)

    response = fast_e2e_harness.client.post(
        "/api/tenancies",
        json={
            "property_id": property_record.id,
            "full_name": "Maya Angelou",
            "email": "maya@example.com",
            "lease_start": "2024-01-01",
            "lease_end": "2025-01-01",
            "monthly_rent": 1200,
            "deposit": 3000,
        },
    )
    assert response.status_code == 200
    payload = response.json()
    tenancy_id = payload["id"]

    list_response = fast_e2e_harness.client.get(
        "/api/tenancies",
        params={"property_id": property_record.id, "active_only": True},
    )
    assert list_response.status_code == 200
    assert list_response.json()["count"] == 1

    show_response = fast_e2e_harness.client.get(f"/api/tenancies/{tenancy_id}")
    assert show_response.status_code == 200
    assert show_response.json()["full_name"] == "Maya Angelou"

    telegram_reply = fast_e2e_harness.orchestrator.handle_telegram_message_with_ui(
        chat_id=77,
        message_text=f"/tenant list {property_record.id}",
    )
    assert "Active tenants" in telegram_reply.text
    assert "Maya Angelou" in telegram_reply.text


@pytest.mark.e2e
def test_tenant_document_upload_persists_memory_and_extracted_fields(fast_e2e_harness: Any) -> None:
    property_record = Property(
        id="prop-tenant-2",
        address="22 Market Street",
        purchase_date=None,
        purchase_price=None,
        current_value=None,
        status=PropertyStatus.OWNED,
    )
    fast_e2e_harness.property_registry.add_property(property_record)

    tenancy = fast_e2e_harness.tenancy_service.create_tenancy(
        property_id=property_record.id,
        full_name="Jane Doe",
        lease_start=date(2024, 1, 1),
        lease_end=date(2025, 1, 1),
        monthly_rent=Decimal("800.00"),
        deposit=Decimal("1600.00"),
    )

    response = fast_e2e_harness.client.post(
        f"/api/tenancies/{tenancy.id}/documents",
        files={"file": ("tenant.txt", b"Tenant: Jane Doe\nEmail: jane@example.com\nPhone: +44 7700 123456\nAnnual Income: 40000", "text/plain")},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["extracted_fields"]["email"] == "jane@example.com"
    assert body["document_type"] in {"text", "txt"}
    assert any(record.payload.tenancy_id == tenancy.id for record in fast_e2e_harness.memory_store.records)


@pytest.mark.e2e
def test_tenant_search_with_date_range_filters_results(fast_e2e_harness: Any) -> None:
    property_record = Property(
        id="prop-tenant-3",
        address="47 Park Avenue",
        purchase_date=None,
        purchase_price=None,
        current_value=None,
        status=PropertyStatus.OWNED,
    )
    fast_e2e_harness.property_registry.add_property(property_record)

    tenancy = fast_e2e_harness.tenancy_service.create_tenancy(
        property_id=property_record.id,
        full_name="Benedict Cumberbatch",
        lease_start=date(2024, 1, 1),
        lease_end=date(2025, 1, 1),
        monthly_rent=Decimal("1500.00"),
        deposit=Decimal("3000.00"),
    )

    document = fast_e2e_harness.tenancy_service.store_document(
        tenancy.id,
        Path(__file__).with_name("fixtures") / "tenant_search.txt",
        filename="tenant_search.txt",
        event_date=date(2025, 2, 1),
    )
    assert document is not None

    matches = fast_e2e_harness.tenancy_service.search_documents(
        query="tenant",
        tenancy_id=tenancy.id,
        date_from=date(2025, 1, 1),
        date_to=date(2025, 3, 1),
    )
    assert len(matches) >= 1


@pytest.mark.e2e
def test_agreement_generation_creates_docx_and_orchestrator_reply(fast_e2e_harness: Any, tmp_path: Path) -> None:
    property_record = Property(
        id="prop-tenant-4",
        address="8 Lake Road",
        purchase_date=None,
        purchase_price=None,
        current_value=None,
        status=PropertyStatus.OWNED,
    )
    fast_e2e_harness.property_registry.add_property(property_record)

    tenancy = fast_e2e_harness.tenancy_service.create_tenancy(
        property_id=property_record.id,
        full_name="Alicia Keys",
        lease_start=date(2024, 1, 1),
        lease_end=date(2025, 1, 1),
        monthly_rent=Decimal("600.00"),
        deposit=Decimal("1200.00"),
    )

    templates_dir = tmp_path / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    template_path = templates_dir / "residential_lease.docx"
    document = Document()
    document.add_paragraph("*|TENANT_FULL_NAME|*\n*|PROPERTY_ADDRESS|*\n*|MONTHLY_RENT|*")
    document.save(template_path)
    fast_e2e_harness.tenancy_service._template_dir = templates_dir
    fast_e2e_harness.tenancy_service._generated_dir = tmp_path / "generated"
    fast_e2e_harness.tenancy_service._generated_dir.mkdir(parents=True, exist_ok=True)

    agreement, unresolved = fast_e2e_harness.tenancy_service.generate_agreement(tenancy.id, template_name="residential_lease")
    assert agreement.stored_path.endswith(".docx")
    assert Path(agreement.stored_path).exists()
    assert unresolved == []

    reply = fast_e2e_harness.orchestrator.handle_telegram_message_with_ui(
        chat_id=99,
        message_text=f"/agreement generate {tenancy.id}",
    )
    assert "Agreement generated successfully" in reply.text
